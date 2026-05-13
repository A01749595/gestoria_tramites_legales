import argparse
import os
import re
from pathlib import Path

from docx import Document
from dotenv import load_dotenv

load_dotenv()

DOCX_PATH = Path(__file__).parent / "tiendas 21" / "TRAMITES.docx"

TIENDAS = [
    ("T-001", "San Antonio Abad",   "CDMX"),
    ("T-003", "Portales",            "CDMX"),
    ("T-007", "Tlalnepantla",        "Tlalnepantla"),
    ("T-008", "Chalco",              "Chalco"),
    ("T-014", "Puebla Outlet",       "Puebla"),
    ("T-015", "Atlixco",             "Atlixco"),
    ("T-016", "Puebla Centro",       "Puebla"),
    ("T-017", "Izúcar",              "Izúcar de Matamoros"),
    ("T-018", "Huamantla",           "Huamantla"),
    ("T-019", "Orizaba",             "Orizaba"),
    ("T-022", "Veracruz",            "Veracruz"),
    ("T-023", "Chilpancingo",        "Chilpancingo"),
    ("T-024", "Zamora",              "Zamora"),
    ("T-025", "Uruapan",             "Uruapan"),
    ("T-028", "Mérida",              "Mérida"),
    ("T-030", "Tuxpan",              "Tuxpan"),
    ("T-032", "Tuxtla Gutiérrez",    "Tuxtla Gutiérrez"),
    ("T-034", "Villahermosa",        "Villahermosa"),
    ("T-035", "Cuautla Bravos",      "Cuautla"),
    ("T-037", "Apizaco",             "Apizaco"),
    ("T-038", "Toluca",              "Toluca"),
]

_NORM = {
    "aviso de funcionamiento":                          "licencia_funcionamiento",
    "uso de suelo":                                     "uso_suelo",
    "anuncio":                                          "anuncio",
    "protección civil":                                 "proteccion_civil",
    "protección civil visto bueno":                     "proteccion_civil",
    "dictamen bomberos":                                "bomberos",
    "convenio de bomberos":                             "bomberos",
    "dictamen de bomberos":                             "bomberos",
    "licencia ambiental":                               "licencia_ambiental",
    "dictamen de medidas contra incendio":              "dictamen_incendio",
    "factibilidad de uso de suelo centro histórico":    "factibilidad_uso_suelo_ch",
    "alineamiento y no oficial":                        "alineamiento",
    "cédula de empadronamiento":                        "cedula_empadronamiento",
    "cedula de empadronamiento":                        "cedula_empadronamiento",
    "aprobación del programa de protección c":          "programa_pc",
    "aprobación del programa de protección civil":      "programa_pc",
    "dictamen de regulación de niveles máximos de sonido": "dictamen_sonido",
}

def normalize(nombre: str) -> str:
    key = nombre.strip().lower()
    for pattern, doc_tipo in _NORM.items():
        if key.startswith(pattern):
            return doc_tipo
    return re.sub(r"[^a-z0-9]+", "_", key).strip("_")


def parse_docx() -> list[dict]:
    doc = Document(DOCX_PATH)

    if len(doc.tables) != len(TIENDAS):
        print(f"ADVERTENCIA: {len(doc.tables)} tablas en docx vs {len(TIENDAS)} tiendas definidas")

    rows = []
    for (tienda_id, tienda_nombre, municipio), table in zip(TIENDAS, doc.tables):
        for row in table.rows[1:]:  # skip header row
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 4:
                continue
            permiso_nombre = cells[1].strip()
            if not permiso_nombre:
                continue
            vigencia_2025 = cells[2].strip() or None
            vigencia_2026 = cells[3].strip() or None
            doc_tipo = normalize(permiso_nombre)
            rows.append({
                "tienda_id":     tienda_id,
                "tienda_nombre": tienda_nombre,
                "municipio":     municipio,
                "doc_tipo":      doc_tipo,
                "permiso_nombre": permiso_nombre,
                "vigencia_2025": vigencia_2025,
                "vigencia_2026": vigencia_2026,
            })
    return rows


def print_table(rows: list[dict]) -> None:
    rows_sorted = sorted(rows, key=lambda r: (r["tienda_id"], r["doc_tipo"]))
    print(f"\n{'Tienda':<10} {'Municipio':<22} {'doc_tipo':<30} {'Vigencia 2025':<18} {'Vigencia 2026'}")
    print("-" * 105)
    for r in rows_sorted:
        print(
            f"{r['tienda_id']:<10} {r['municipio']:<22} {r['doc_tipo']:<30} "
            f"{(r['vigencia_2025'] or ''):<18} {r['vigencia_2026'] or ''}"
        )
    print(f"\nTotal: {len(rows)} registros")


def upsert_db(rows: list[dict]) -> None:
    from supabase import create_client
    url = os.environ["SUPABASE_URL"]
    key = os.environ["SUPABASE_KEY"]
    sb = create_client(url, key)

    upserted = 0
    for r in rows:
        sb.table("tienda_permisos").upsert(r, on_conflict="tienda_id,doc_tipo").execute()
        upserted += 1
    print(f"Upserted {upserted} registros en tienda_permisos.")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--dry-run", action="store_true", help="Solo imprime, no escribe en DB")
    args = ap.parse_args()

    rows = parse_docx()
    print_table(rows)

    if not args.dry_run:
        upsert_db(rows)
    else:
        print("\n[dry-run] No se escribió en la base de datos.")


if __name__ == "__main__":
    main()
